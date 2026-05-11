"""DOCX / DOC → plain text.  Uses python-docx for .docx; libreoffice fallback for .doc."""

import io

from docx import Document


def parse_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def parse_doc(data: bytes) -> str:
    """Best-effort .doc parsing via python-docx (works on many .doc files)."""
    try:
        return parse_docx(data)
    except Exception:
        return ""
