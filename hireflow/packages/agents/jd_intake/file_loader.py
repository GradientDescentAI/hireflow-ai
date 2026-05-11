"""
Load raw text from the various JD source types accepted by JD-001.

source_type: 'text' | 'pdf' | 'docx' | 'url'

PDF text extraction uses pdfplumber; falls back to pytesseract OCR for
scanned/image-only PDFs.  URL extraction strips HTML tags.
"""

import io
import re

import httpx
import pdfplumber
from docx import Document


_URL_TIMEOUT = 30


def load_text(source: str | bytes, source_type: str) -> str:
    """Return plain text extracted from source.

    Args:
        source: raw bytes (pdf/docx) or string (text/url).
        source_type: one of 'text', 'pdf', 'docx', 'url'.

    Returns:
        Extracted text, stripped of leading/trailing whitespace.

    Raises:
        ValueError: unrecognised source_type.
        httpx.HTTPError: network error on 'url' source.
    """
    if source_type == "text":
        text = source if isinstance(source, str) else source.decode("utf-8", errors="replace")
        return text.strip()

    if source_type == "pdf":
        return _load_pdf(source if isinstance(source, bytes) else open(source, "rb").read())

    if source_type == "docx":
        return _load_docx(source if isinstance(source, bytes) else open(source, "rb").read())

    if source_type == "url":
        return _load_url(source if isinstance(source, str) else source.decode())

    raise ValueError(f"Unknown source_type: {source_type!r}. Valid: text, pdf, docx, url")


def _load_pdf(data: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    text = "\n".join(text_parts).strip()

    if not text:
        # Scanned PDF — fall back to OCR
        text = _ocr_pdf(data)

    return text


def _ocr_pdf(data: bytes) -> str:
    try:
        import pytesseract
        from pdf2image import convert_from_bytes  # type: ignore[import]
    except ImportError:
        return ""

    images = convert_from_bytes(data)
    return "\n".join(pytesseract.image_to_string(img) for img in images).strip()


def _load_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract table cell text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def _load_url(url: str) -> str:
    resp = httpx.get(url, follow_redirects=True, timeout=_URL_TIMEOUT)
    resp.raise_for_status()
    # Strip HTML tags — keep text content
    text = re.sub(r"<style[^>]*>.*?</style>", " ", resp.text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<script[^>]*>.*?</script>", " ", text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"&[a-z]+;", " ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()
